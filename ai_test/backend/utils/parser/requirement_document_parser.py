"""
需求文档解析器 - 支持从多种格式文档中提取文本内容和嵌入图片
支持格式：PDF、DOCX、TXT、MD
"""
import os
import io
import base64
import logging

logger = logging.getLogger(__name__)

MAX_IMAGES = 15
MAX_IMAGE_SIZE = 5 * 1024 * 1024


def extract_text_from_pdf(file_content: bytes) -> str:
    """从PDF文件中提取文本"""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_content: bytes) -> str:
    """从DOCX文件中提取文本"""
    from docx import Document
    doc = Document(io.BytesIO(file_content))
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def extract_text_from_txt(file_content: bytes) -> str:
    """从TXT/MD文件中提取文本"""
    # 尝试多种编码
    for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
        try:
            return file_content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return file_content.decode('utf-8', errors='ignore')


def extract_text_from_url(url: str) -> str:
    """从URL中提取文本内容"""
    import requests
    from bs4 import BeautifulSoup

    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'
    }
    response = requests.get(url, headers=headers, timeout=30)
    response.raise_for_status()

    content_type = response.headers.get('content-type', '').lower()

    # 如果是PDF
    if 'application/pdf' in content_type:
        return extract_text_from_pdf(response.content)

    # 如果是Word文档
    if 'application/vnd.openxmlformats-officedocument' in content_type or \
       'application/msword' in content_type:
        return extract_text_from_docx(response.content)

    # 默认当作HTML处理
    soup = BeautifulSoup(response.text, 'html.parser')

    # 移除script和style标签
    for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
        tag.decompose()

    # 提取主要内容
    # 优先查找 article、main 等语义化标签
    main_content = soup.find('article') or soup.find('main') or soup.find('body')
    if main_content:
        text = main_content.get_text(separator='\n', strip=True)
    else:
        text = soup.get_text(separator='\n', strip=True)

    # 清理多余空行
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    return '\n'.join(lines)


def extract_images_from_docx(file_content: bytes) -> list[tuple[str, str]]:
    """从 DOCX 文件中提取所有嵌入图片，返回 [(base64_data, content_type)]"""
    from docx import Document
    doc = Document(io.BytesIO(file_content))
    images = []
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    blob = rel.target_part.blob
                    if len(blob) > MAX_IMAGE_SIZE:
                        continue
                    ct = rel.target_part.content_type or "image/png"
                    b64 = base64.b64encode(blob).decode('utf-8')
                    images.append((b64, ct))
                    if len(images) >= MAX_IMAGES:
                        break
                except Exception:
                    continue
    except Exception as e:
        logger.warning(f"DOCX图片提取失败: {e}")
    return images


def extract_images_from_pdf(file_content: bytes) -> list[tuple[str, str]]:
    """从 PDF 文件中提取嵌入图片（使用 pdfplumber 获取图片区域截图）"""
    images = []
    try:
        import pdfplumber
        from PIL import Image
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                if len(images) >= MAX_IMAGES:
                    break
                page_images = page.images
                if not page_images:
                    continue
                pil_img = page.to_image(resolution=150).original
                for img_info in page_images:
                    if len(images) >= MAX_IMAGES:
                        break
                    try:
                        x0, top, x1, bottom = img_info["x0"], img_info["top"], img_info["x1"], img_info["bottom"]
                        cropped = pil_img.crop((x0, top, x1, bottom))
                        if cropped.width < 50 or cropped.height < 50:
                            continue
                        buf = io.BytesIO()
                        cropped.save(buf, format="PNG")
                        blob = buf.getvalue()
                        if len(blob) > MAX_IMAGE_SIZE:
                            continue
                        b64 = base64.b64encode(blob).decode('utf-8')
                        images.append((b64, "image/png"))
                    except Exception:
                        continue
    except ImportError:
        logger.warning("PDF图片提取需要 Pillow 库")
    except Exception as e:
        logger.warning(f"PDF图片提取失败: {e}")
    return images


def extract_text_and_images_from_file(
    filename: str, file_content: bytes
) -> tuple[str, list[tuple[str, str]]]:
    """
    提取文本 + 嵌入图片。
    返回 (text, [(base64_data, content_type)])
    """
    text = extract_text_from_file(filename, file_content)
    images = []
    ext = os.path.splitext(filename)[1].lower()
    try:
        if ext in ('.docx', '.doc'):
            images = extract_images_from_docx(file_content)
        elif ext == '.pdf':
            images = extract_images_from_pdf(file_content)
    except Exception as e:
        logger.warning(f"图片提取失败（不影响文本提取）: {e}")
    return text, images


def extract_text_from_file(filename: str, file_content: bytes) -> str:
    """
    根据文件类型自动选择解析器提取文本

    Args:
        filename: 文件名
        file_content: 文件二进制内容

    Returns:
        提取的文本内容
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == '.pdf':
        return extract_text_from_pdf(file_content)
    elif ext in ('.docx', '.doc'):
        return extract_text_from_docx(file_content)
    elif ext in ('.txt', '.md', '.markdown', '.rst'):
        return extract_text_from_txt(file_content)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，支持的格式为: PDF, DOCX, TXT, MD")


# 支持的文件类型
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.md', '.markdown'}
SUPPORTED_MIME_TYPES = {
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'application/msword',
    'text/plain',
    'text/markdown',
}

# 文件大小限制 (10MB)
MAX_FILE_SIZE = 10 * 1024 * 1024
